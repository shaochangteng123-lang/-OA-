#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
财务区上线计划文档生成脚本
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run)
        run.font.bold = True
    return heading

def add_paragraph(doc, text, bold=False, font_size=12):
    """添加段落"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        set_chinese_font(run)
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    return para

def create_onboarding_plan():
    """创建上线计划文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    title = doc.add_heading('YuliLog 工作日志系统 - 财务区上线计划', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run)
        run.font.size = Pt(18)
        run.font.bold = True

    # 基本信息
    info_para = doc.add_paragraph()
    info_para.add_run('上线时间：').bold = True
    info_para.add_run('2026年4月4日\n')
    info_para.add_run('培训时间：').bold = True
    info_para.add_run('2026年4月4日 上午10:00\n')
    info_para.add_run('培训方式：').bold = True
    info_para.add_run('线上会议（微信群通知）\n')
    info_para.add_run('编制人：').bold = True
    info_para.add_run('羽隶\n')
    info_para.add_run('编制日期：').bold = True
    info_para.add_run(f'{datetime.date.today().strftime("%Y年%m月%d日")}\n')
    for run in info_para.runs:
        set_chinese_font(run)
        run.font.size = Pt(11)

    doc.add_paragraph()  # 空行

    # ========== 第一部分：上线概述 ==========
    add_heading(doc, '一、上线概述', level=1)

    add_paragraph(doc, '1.1 上线背景', bold=True)
    add_paragraph(doc,
        'YuliLog 工作日志系统财务区模块已完成开发和测试，具备正式上线条件。'
        '该模块旨在规范公司报销流程，提高财务管理效率，实现报销申请、审批、付款、'
        '确认收款的全流程数字化管理。')

    doc.add_paragraph()
    add_paragraph(doc, '1.2 上线范围', bold=True)
    add_paragraph(doc, '本次上线包含以下功能模块：')

    features = [
        '基础报销管理（< 1000元）',
        '大额报销管理（≥ 1000元）',
        '商务报销管理（不限金额）',
        '报销审批流程（管理员审批、总经理审批）',
        '付款管理（单笔付款、批量付款）',
        '付款凭证OCR识别与验证',
        '发票管理（上传、查看、下载）',
        '核减发票管理（自动核减、手动核减）',
        '报销统计与查询',
        '收款确认功能'
    ]

    for feature in features:
        para = doc.add_paragraph(feature, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)

    doc.add_paragraph()
    add_paragraph(doc, '1.3 上线时间安排', bold=True)

    schedule_table = doc.add_table(rows=5, cols=3)
    schedule_table.style = 'Light Grid Accent 1'

    schedule_data = [
        ['时间', '事项', '负责人'],
        ['2026-04-03 下午', '发送微信群通知，发放操作手册', '羽隶'],
        ['2026-04-04 上午10:00', '线上培训会议', '羽隶'],
        ['2026-04-04 下午', '系统正式上线，开放使用', '羽隶'],
        ['2026-04-05 起', '持续跟进，收集反馈，优化改进', '羽隶']
    ]

    for i, row_data in enumerate(schedule_data):
        row = schedule_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.size = Pt(11)
                    if i == 0:  # 表头加粗
                        run.font.bold = True

    doc.add_page_break()

    # ========== 第二部分：操作手册 ==========
    add_heading(doc, '二、财务区操作手册', level=1)

    add_heading(doc, '2.1 系统访问与登录', level=2)

    add_paragraph(doc, '（1）系统访问地址', bold=True)
    access_table = doc.add_table(rows=3, cols=3)
    access_table.style = 'Light Grid Accent 1'
    access_data = [
        ['阶段', '访问地址', '适用条件'],
        ['当前内网阶段', 'http://192.168.100.7:8899', '需连接公司 Wi-Fi 或办公室内网'],
        ['正式上线后', 'https://yuliseek.com', '任意网络均可访问'],
    ]
    for i, row_data in enumerate(access_data):
        row = access_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.size = Pt(11)
                    if i == 0:
                        run.font.bold = True

    doc.add_paragraph()
    add_paragraph(doc,
        '⚠️  注意：内网阶段请确保手机或电脑已连接公司 Wi-Fi，'
        '否则无法访问系统。推荐使用 Chrome 或 Edge 浏览器。')

    doc.add_paragraph()
    add_paragraph(doc, '（2）账号获取方式', bold=True)
    add_paragraph(doc,
        '系统账号由管理员统一创建，员工不能自行注册。'
        '账号开通后，管理员会将您的用户名和初始密码单独告知您。')

    account_tips = [
        '用户名：由管理员分配，一般为您的姓名拼音或工号',
        '初始密码：由管理员告知，长度至少6位',
        '首次登录后建议立即修改密码（在系统右上角个人设置中修改）',
        '账号遗忘或密码忘记，请联系管理员重置'
    ]
    for item in account_tips:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '（3）登录步骤', bold=True)

    login_steps = [
        '打开浏览器（推荐 Chrome 或 Edge），在地址栏输入访问地址，回车',
        '进入登录页面，在"用户名"输入框填写您的账号',
        '在"密码"输入框填写密码（输入时密码不显示，属正常现象）',
        '点击"登录"按钮，或直接按键盘 Enter 键',
        '登录成功后，自动跳转到系统首页'
    ]
    for i, step in enumerate(login_steps, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '【截图占位：登录页面截图】', bold=True)
    doc.add_paragraph('（此处插入登录页面截图）')

    doc.add_paragraph()
    add_paragraph(doc, '（4）常见登录问题', bold=True)
    login_faq = [
        ('提示"用户名或密码错误"', '请确认用户名和密码是否正确，注意区分大小写'),
        ('页面打不开', '请检查是否已连接公司 Wi-Fi（内网阶段），或尝试刷新页面'),
        ('密码忘记', '联系管理员重置密码，管理员会告知新的初始密码'),
        ('登录后自动退出', '系统有登录超时机制，长时间不操作会自动退出，重新登录即可'),
    ]
    for q, a in login_faq:
        para = doc.add_paragraph()
        run_q = para.add_run(f'问：{q}\n')
        run_q.bold = True
        set_chinese_font(run_q)
        run_q.font.size = Pt(11)
        run_a = para.add_run(f'答：{a}')
        set_chinese_font(run_a)
        run_a.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_page_break()

    add_heading(doc, '2.2 系统概述', level=2)
    add_paragraph(doc,
        '财务区是 YuliLog 工作日志系统的核心模块之一，用于管理公司的报销业务。'
        '系统支持三种报销类型，实现了从申请、审批、付款到确认收款的完整流程闭环。')

    doc.add_paragraph()
    add_heading(doc, '2.2 报销类型说明', level=2)

    # 基础报销
    add_paragraph(doc, '（1）基础报销（< 1000元）', bold=True)
    add_paragraph(doc, '适用范围：金额小于1000元的日常报销')
    add_paragraph(doc, '报销品类包括：')

    basic_categories = [
        '行政办公类（9种）：办公设备、办公器具、办公耗材、劳保用品、水电费、物业费、清洁费、差旅费、茶水费',
        '三方服务类（8种）：图文服务、快递服务、车辆服务、物流服务、公证服务、咨询服务、技术服务、制证服务',
        '业务提升类（3种）：培训费、知识付费、专家指导'
    ]

    for cat in basic_categories:
        para = doc.add_paragraph(cat, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    add_paragraph(doc, '审批流程：员工提交 → 管理员审批 → 付款 → 员工确认收款 → 完成')

    doc.add_paragraph()

    # 大额报销
    add_paragraph(doc, '（2）大额报销（≥ 1000元）', bold=True)
    add_paragraph(doc, '适用范围：金额大于等于1000元的报销')
    add_paragraph(doc, '报销品类：与基础报销相同')
    add_paragraph(doc, '审批流程：员工提交 → 管理员审批 → 付款 → 员工确认收款 → 完成')

    doc.add_paragraph()

    # 商务报销
    add_paragraph(doc, '（3）商务报销（不限金额）', bold=True)
    add_paragraph(doc, '适用范围：商务活动相关的报销，不限金额')
    add_paragraph(doc, '报销品类：商务接待、客户拜访、商务差旅、会议活动')
    add_paragraph(doc, '审批流程：员工提交 → 总经理审批 → 管理员付款 → 员工确认收款 → 完成')

    doc.add_paragraph()
    add_heading(doc, '2.3 特殊规则说明', level=2)

    add_paragraph(doc, '（1）运输类发票月度额度限制', bold=True)
    add_paragraph(doc,
        '运输、交通、汽油、柴油类发票每人每月最多报销1500元。'
        '超出部分系统将自动核减，核减金额按比例分配到每张运输类发票。')

    doc.add_paragraph()
    add_paragraph(doc, '（2）核减发票', bold=True)
    add_paragraph(doc,
        '员工可上传核减发票作为凭证，这些发票不计入报销金额。'
        '核减发票用于存档不能报销的发票。')

    doc.add_paragraph()
    add_paragraph(doc, '（3）付款凭证验证', bold=True)
    add_paragraph(doc,
        '管理员上传付款回单时，系统会自动OCR识别收款人、账号、金额等信息，'
        '并与报销单信息进行验证，确保付款准确无误。')

    doc.add_page_break()

    # ========== 员工操作指南 ==========
    add_heading(doc, '2.4 员工操作指南', level=2)

    add_paragraph(doc, '步骤1：创建报销单', bold=True)
    steps_1 = [
        '登录系统后，点击左侧菜单"财务区"，选择对应的报销类型（基础报销/大额报销/商务报销）',
        '点击"创建报销单"按钮',
        '填写报销信息：报销标题、报销类别、报销月份、服务对象等',
        '商务报销需额外填写：商务类型、客户名称',
        '点击"保存草稿"可暂存，点击"提交审批"正式提交'
    ]
    for i, step in enumerate(steps_1, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '步骤2：上传发票', bold=True)
    steps_2 = [
        '在报销单详情页，点击"上传发票"按钮',
        '选择发票PDF文件（支持多张发票）',
        '系统会自动识别发票信息（金额、日期、发票号码等）',
        '如有无票收据，点击"上传无票收据"',
        '如有核减发票，点击"上传核减发票"（不计入报销金额）',
        '确认发票信息无误后，点击"提交审批"'
    ]
    for i, step in enumerate(steps_2, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '步骤3：等待审批', bold=True)
    add_paragraph(doc,
        '提交后，报销单进入待审批状态。基础/大额报销由管理员审批，'
        '商务报销由总经理审批。审批通过后，报销单进入待付款状态。')

    doc.add_paragraph()
    add_paragraph(doc, '步骤4：确认收款', bold=True)
    add_paragraph(doc,
        '管理员完成付款并上传付款回单后，员工会收到通知。'
        '员工确认收到款项后，点击"确认收款"按钮，报销流程完成。')

    doc.add_paragraph()
    add_paragraph(doc, '步骤5：查看报销历史', bold=True)
    add_paragraph(doc,
        '在报销列表页可查看所有报销单的状态和历史记录。'
        '点击报销单可查看详情、下载发票、查看付款回单等。')

    doc.add_page_break()

    # ========== 管理员操作指南 ==========
    add_heading(doc, '2.5 管理员操作指南', level=2)

    add_paragraph(doc, '审批管理', bold=True)
    admin_steps_1 = [
        '进入"审批中心"（管理员视图），查看所有待审批的基础报销和大额报销',
        '点击报销单查看详情，核对发票信息',
        '审批通过：点击"通过"按钮',
        '审批驳回：点击"驳回"按钮，填写驳回原因'
    ]
    for i, step in enumerate(admin_steps_1, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '付款操作（单笔）', bold=True)
    admin_steps_2 = [
        '审批通过后，在"已批准待付款"列表中找到该报销单',
        '点击"发起付款"，确认付款金额和收款信息',
        '完成银行转账后，上传付款回单',
        '系统自动OCR验证，验证通过后确认上传',
        '员工确认收款后，该报销单完成'
    ]
    for i, step in enumerate(admin_steps_2, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '付款操作（批量）', bold=True)
    admin_steps_3 = [
        '在"已批准待付款"列表中，勾选多个同一收款人、同一报销类型的报销单',
        '点击"批量付款"，系统自动合并金额',
        '完成银行转账后，上传付款回单',
        '系统OCR验证通过后批量确认'
    ]
    for i, step in enumerate(admin_steps_3, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_heading(doc, '2.6 总经理操作指南', level=2)

    add_paragraph(doc, '审批商务报销', bold=True)
    gm_steps = [
        '进入"总经理审批中心"，查看所有待审批的商务报销',
        '点击报销单查看详情（商务类型、客户名称、发票等）',
        '审批通过：点击"通过"按钮',
        '审批驳回：点击"驳回"按钮，填写驳回原因',
        '通过后，报销单转由管理员处理付款'
    ]
    for i, step in enumerate(gm_steps, 1):
        para = doc.add_paragraph(f'{i}. {step}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 报销状态说明 ==========
    add_heading(doc, '2.7 报销状态说明', level=2)

    status_table = doc.add_table(rows=8, cols=2)
    status_table.style = 'Light Grid Accent 1'

    status_data = [
        ['状态', '说明'],
        ['草稿（draft）', '报销单已创建但未提交，可继续修改'],
        ['待审批（pending）', '已提交，等待审批'],
        ['已审批（approved）', '审批通过，等待管理员付款'],
        ['待确认收款（payment_uploaded）', '管理员已上传付款回单，等待员工确认收款'],
        ['已完成（completed）', '员工确认收款，流程结束'],
        ['已驳回（rejected）', '审批不通过，员工可修改后重新提交'],
        ['已撤回', '员工主动撤回报销单（可重新修改提交）']
    ]

    for i, row_data in enumerate(status_data):
        row = status_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.size = Pt(11)
                    if i == 0:
                        run.font.bold = True

    doc.add_paragraph()
    doc.add_page_break()

    # ========== 第三部分：培训计划 ==========
    add_heading(doc, '三、培训计划', level=1)

    add_heading(doc, '3.1 培训概述', level=2)
    train_info = doc.add_paragraph()
    train_info.add_run('培训时间：').bold = True
    train_info.add_run('2026年4月4日 上午10:00-11:30\n')
    train_info.add_run('培训方式：').bold = True
    train_info.add_run('线上视频培训（微信视频会议）\n')
    train_info.add_run('培训对象：').bold = True
    train_info.add_run('全体员工（含管理员、总经理）\n')
    train_info.add_run('培训讲师：').bold = True
    train_info.add_run('羽隶\n')
    for run in train_info.runs:
        set_chinese_font(run)
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_heading(doc, '3.2 培训内容大纲', level=2)

    add_paragraph(doc, '一、系统简介（5分钟）', bold=True)
    intro_items = [
        '财务区模块功能概述',
        '报销数字化的意义和优势',
        '系统访问地址和登录方式'
    ]
    for item in intro_items:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '二、报销类型介绍（10分钟）', bold=True)
    type_items = [
        '基础报销：适用场景、金额范围、报销品类',
        '大额报销：与基础报销的区别',
        '商务报销：审批流程的不同点',
        '如何选择正确的报销类型'
    ]
    for item in type_items:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '三、员工操作演示（25分钟）', bold=True)
    employee_demo = [
        '创建报销单：演示填写流程',
        '上传发票：演示OCR识别效果',
        '上传无票收据、核减发票的操作方法',
        '提交审批：确认信息并提交',
        '确认收款：演示确认操作',
        '查看报销历史和状态'
    ]
    for item in employee_demo:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '四、管理员操作演示（15分钟）', bold=True)
    admin_demo = [
        '审批中心：查看和处理待审批报销单',
        '审批通过/驳回操作演示',
        '发起付款：单笔和批量付款演示',
        '上传付款回单：OCR验证演示',
        '报销统计查询功能'
    ]
    for item in admin_demo:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '五、总经理操作演示（5分钟）', bold=True)
    gm_demo = [
        '进入总经理审批中心',
        '查看商务报销详情',
        '审批通过/驳回操作'
    ]
    for item in gm_demo:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '六、特殊规则说明（10分钟）', bold=True)
    special_rules = [
        '运输类发票月度1500元额度限制说明',
        '核减发票的使用场景和上传方法',
        '付款凭证OCR验证的注意事项',
        '报销单驳回后的修改和重提流程'
    ]
    for item in special_rules:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '七、问答环节（20分钟）', bold=True)
    qa_items = [
        '解答员工疑问',
        '演示补充操作（根据提问）',
        '收集使用建议'
    ]
    for item in qa_items:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_heading(doc, '3.3 培训注意事项', level=2)
    notes = [
        '培训前请确保所有参与者已能正常登录系统',
        '请准备好公司报销账号信息（用于演示）',
        '管理员和总经理需分别演示各自视角的操作',
        '培训结束后，将操作手册分发给所有员工留存',
        '培训后设立问题反馈渠道，及时响应员工使用问题'
    ]
    for i, note in enumerate(notes, 1):
        para = doc.add_paragraph(f'{i}. {note}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 第四部分：微信群通知话术 ==========
    add_heading(doc, '四、微信群通知话术', level=1)

    add_paragraph(doc, '【发送时间】2026年4月3日（今日）下午', bold=True)
    doc.add_paragraph()

    add_paragraph(doc, '以下为发送至微信群的通知内容：', bold=False)
    doc.add_paragraph()

    # 通知框
    notice_para = doc.add_paragraph()
    notice_text = (
        '📢 OA系统上线通知\n\n'
        '各位同事：\n\n'
        '大家好！\n\n'
        '公司 OA 工作日志系统「财务区」模块已完成开发和测试，将于本周五（4月4日）'
        '正式上线启用。财务区实现了报销申请、审批、付款、收款确认的全流程数字化管理，'
        '将极大提升我们的报销工作效率。\n\n'
        '📋 主要功能：\n'
        '✅ 基础报销（< 1000元）\n'
        '✅ 大额报销（≥ 1000元）\n'
        '✅ 商务报销（接待/拜访/差旅/会议）\n'
        '✅ 发票上传与管理\n'
        '✅ 在线审批与付款确认\n\n'
        '📅 培训安排：\n'
        '时间：2026年4月4日（明天）上午 10:00\n'
        '方式：微信视频会议（届时发送入会链接）\n'
        '时长：约1.5小时\n'
        '请所有同事准时参加！\n\n'
        '📄 操作手册已随本通知发出，请大家提前熟悉。\n\n'
        '如有问题，请随时联系羽隶。\n\n'
        '北京羽隶工程咨询有限公司\n'
        '2026年4月3日'
    )
    notice_run = notice_para.add_run(notice_text)
    set_chinese_font(notice_run)
    notice_run.font.size = Pt(11)

    doc.add_paragraph()
    add_paragraph(doc, '【注意事项】', bold=True)
    notice_notes = [
        '发送通知时，同步发送本操作手册（PDF格式）',
        '通知发出后，在微信群内@所有人，确保覆盖',
        '次日（4月4日）上午9:30，发送培训会议链接',
        '培训前再次提醒参会'
    ]
    for i, note in enumerate(notice_notes, 1):
        para = doc.add_paragraph(f'{i}. {note}')
        for run in para.runs:
            set_chinese_font(run)
            run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 培训会议通知话术 ==========
    add_heading(doc, '五、培训会议通知话术', level=1)

    add_paragraph(doc, '【发送时间】2026年4月4日（培训当天）上午9:30', bold=True)
    doc.add_paragraph()

    meeting_notice = (
        '📢 培训会议提醒\n\n'
        '各位同事：\n\n'
        '财务区上线培训将于今日上午10:00开始，请准时参加！\n\n'
        '📅 培训时间：今天（4月4日）上午10:00\n'
        '⏱ 预计时长：1.5小时\n\n'
        '👇 请点击以下链接进入微信视频会议：\n'
        '[会议链接请在此处填写]\n\n'
        '请提前准备好：\n'
        '✅ 登录OA系统账号（确认能正常登录）\n'
        '✅ 查看昨日发送的操作手册\n'
        '✅ 准备好想问的问题\n\n'
        '如临时有事不能参加，请提前告知。\n\n'
        '羽隶\n'
        '2026年4月4日'
    )

    meeting_para = doc.add_paragraph()
    meeting_run = meeting_para.add_run(meeting_notice)
    set_chinese_font(meeting_run)
    meeting_run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 第五部分：风险与应急预案 ==========
    add_heading(doc, '六、风险评估与应急预案', level=1)

    add_paragraph(doc, '6.1 常见风险', bold=True)

    risk_table = doc.add_table(rows=5, cols=3)
    risk_table.style = 'Light Grid Accent 1'

    risk_data = [
        ['风险项', '可能影响', '应对措施'],
        ['员工不熟悉操作', '报销提交错误', '培训演示 + 操作手册 + 一对一指导'],
        ['发票上传失败', '无法提交报销', '检查文件格式（支持PDF/PNG/JPG），文件大小不超过10MB'],
        ['付款回单验证不通过', '付款流程阻断', '检查账号信息是否与报销单一致，联系管理员重新上传'],
        ['运输类额度不足', '自动核减金额', '系统自动处理，员工知悉核减规则即可']
    ]

    for i, row_data in enumerate(risk_data):
        row = risk_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

    doc.add_paragraph()
    add_paragraph(doc, '6.2 技术支持联系方式', bold=True)
    add_paragraph(doc, '系统使用问题请联系：羽隶（微信/电话）')
    add_paragraph(doc, '反馈问题请在微信群中说明，我们将在1个工作日内响应。')

    doc.add_paragraph()
    doc.add_page_break()

    # ========== 附录 ==========
    add_heading(doc, '附录：名词解释', level=1)

    glossary_table = doc.add_table(rows=9, cols=2)
    glossary_table.style = 'Light Grid Accent 1'

    glossary_data = [
        ['术语', '说明'],
        ['基础报销', '金额小于1000元的日常费用报销，由管理员审批'],
        ['大额报销', '金额大于等于1000元的费用报销，由管理员审批'],
        ['商务报销', '商务接待、拜访、差旅、会议费用，由总经理审批'],
        ['核减发票', '不计入报销金额的发票，仅作凭证存档'],
        ['运输类额度', '每人每月最多报销运输/交通/汽油/柴油类发票1500元'],
        ['付款回单', '银行转账完成后的付款凭证截图，用于确认付款'],
        ['OCR识别', '系统自动识别付款回单中的收款人、账号、金额等信息'],
        ['审批中心', '管理员处理报销审批和付款的操作界面']
    ]

    for i, row_data in enumerate(glossary_data):
        row = glossary_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.size = Pt(11)
                    if i == 0:
                        run.font.bold = True

    doc.add_paragraph()

    # 保存
    output_path = '/Volumes/YULi-SCT/北京羽隶工程咨询有限公司/OA办公系统开发中/上线计划.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    create_onboarding_plan()
