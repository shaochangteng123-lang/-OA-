#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
财务区用户操作手册生成脚本
图文并茂版本，包含目录
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ========== 路径配置 ==========
IMG_DIR = '/Volumes/YULi-SCT/北京羽隶工程咨询有限公司/OA办公系统开发/计划/网站截图/'
OUTPUT_PATH = '/Volumes/YULi-SCT/北京羽隶工程咨询有限公司/OA办公系统开发/财务区用户操作手册.docx'

# 截图文件映射（按时间顺序对应页面）
IMGS = {
    'login':           IMG_DIR + 'ScreenShot_2026-04-07_093924_822.jpg',
    'home':            IMG_DIR + 'ScreenShot_2026-04-07_094000_096.jpg',
    'onboarding':      IMG_DIR + 'ScreenShot_2026-04-07_094015_293.png',
    'reimb_list':      IMG_DIR + 'ScreenShot_2026-04-07_094035_391.png',
    'new_basic':       IMG_DIR + 'ScreenShot_2026-04-07_094041_779.png',
    'ocr_result':      IMG_DIR + 'ScreenShot_2026-04-07_094049_317.png',
    'large_filled':    IMG_DIR + 'ScreenShot_2026-04-07_094056_162.png',
    'business_filled': IMG_DIR + 'ScreenShot_2026-04-07_094106_036.png',
    'large_empty':     IMG_DIR + 'ScreenShot_2026-04-07_094115_713.png',
    'business_empty':  IMG_DIR + 'ScreenShot_2026-04-07_094124_321.png',
    'auto_deduction':  IMG_DIR + 'ScreenShot_2026-04-07_094130_885.png',
    'manual_deduction':IMG_DIR + 'ScreenShot_2026-04-07_094142_719.png',
    'submitted_list':  IMG_DIR + 'ScreenShot_2026-04-07_094148_961.png',
    'rejected':        IMG_DIR + 'ScreenShot_2026-04-07_094155_290.png',
    'pay_uploaded':    IMG_DIR + 'ScreenShot_2026-04-07_094201_859.png',
    'completed_list':  IMG_DIR + 'ScreenShot_2026-04-07_094209_130.png',
    'admin_approval':  IMG_DIR + 'ScreenShot_2026-04-07_094218_413.png',
    'admin_flow':      IMG_DIR + 'ScreenShot_2026-04-07_094226_828.png',
    'admin_query':     IMG_DIR + 'ScreenShot_2026-04-07_094233_923.png',
    'admin_invoice':   IMG_DIR + 'ScreenShot_2026-04-07_094243_809.png',
    'admin_pending':   IMG_DIR + 'ScreenShot_2026-04-07_094253_920.png',
    'admin_pay_detail':IMG_DIR + 'ScreenShot_2026-04-07_094301_414.png',
    'upload_proof':    IMG_DIR + '0ce89dfcca495da43f719e6c40232c12.png',
    'proof_done':      IMG_DIR + '56e9e06a5f7a61c6d5b1fae4cdc5b2cb.png',
    'gm_home':         IMG_DIR + '7cdb2468db7fcce43bac2632fa13f464.png',
    'gm_approval':     IMG_DIR + 'c94f32fcc4c4a2cff8d4fc41dd1128dc.png',
    'gm_detail':       IMG_DIR + 'de186876e646a8c91c742adecab1240a.png',
}

# ========== 工具函数 ==========

def set_font(run, size=11, bold=False, color=None):
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=22, bold=True)
    return p

def add_h1(doc, text):
    h = doc.add_heading('', level=1)
    h.clear()
    run = h.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    return h

def add_h2(doc, text):
    h = doc.add_heading('', level=2)
    h.clear()
    run = h.add_run(text)
    set_font(run, size=13, bold=True, color=(68, 114, 196))
    return h

def add_h3(doc, text):
    h = doc.add_heading('', level=3)
    h.clear()
    run = h.add_run(text)
    set_font(run, size=12, bold=True, color=(84, 130, 53))
    return h

def add_para(doc, text, size=11, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_step(doc, number, text):
    """添加步骤段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_num = p.add_run(f'步骤{number}：')
    set_font(run_num, size=11, bold=True, color=(192, 80, 77))
    run_text = p.add_run(text)
    set_font(run_text, size=11)
    return p

def add_tip(doc, text):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'💡 提示：{text}')
    set_font(run, size=10, color=(128, 100, 162))
    return p

def add_warning(doc, text):
    """添加注意框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'⚠️  注意：{text}')
    set_font(run, size=10, bold=True, color=(192, 80, 77))
    return p

def add_bullet(doc, text, indent_level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent_level * 0.75)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_img(doc, key, caption=None, width=6.0):
    path = IMGS.get(key)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    else:
        p = doc.add_paragraph(f'[图片：{key}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            set_font(run, size=9, color=(128, 128, 128))
    doc.add_paragraph()
    return p

def add_divider(doc):
    doc.add_paragraph()

def add_status_table(doc):
    """添加报销状态说明表"""
    data = [
        ['状态名称', '含义', '下一步操作'],
        ['草稿', '已创建但未提交，可继续修改', '完善信息后点击"提交审批"'],
        ['待审批', '已提交，等待审批人处理', '等待审批结果通知'],
        ['已驳回', '审批不通过，需修改后重新提交', '查看驳回原因，修改后重新提交'],
        ['已审批', '审批通过，等待财务付款', '等待财务处理付款'],
        ['待确认收款', '财务已付款并上传回单，等待确认', '确认收到款项后点击"确认收款"'],
        ['已完成', '全流程结束', '可查看历史记录'],
    ]
    table = doc.add_table(rows=len(data), cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=10, bold=(i == 0))
    doc.add_paragraph()

# ========== 主文档生成函数 ==========

def create_manual():
    doc = Document()

    # 页面设置：A4
    from docx.shared import Cm
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(doc, 'YuliLog 工作日志系统')
    add_title(doc, '财务区用户操作手册')
    doc.add_paragraph()
    p = doc.add_paragraph('北京羽隶工程咨询有限公司')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs: set_font(run, size=13)
    p2 = doc.add_paragraph('2026年4月')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs: set_font(run, size=12, color=(128,128,128))
    doc.add_page_break()

    # ===== 目录（手动） =====
    add_h1(doc, '目  录')
    toc_items = [
        ('一、系统简介', '3'),
        ('二、登录与账号设置', '3'),
        ('    2.1  系统访问地址', '3'),
        ('    2.2  登录步骤', '3'),
        ('    2.3  完善个人信息（必做）', '4'),
        ('三、员工操作指南', '5'),
        ('    3.1  报销类型说明', '5'),
        ('    3.2  创建基础报销单', '5'),
        ('    3.3  创建大额报销单', '7'),
        ('    3.4  创建商务报销单', '7'),
        ('    3.5  上传发票与OCR识别', '8'),
        ('    3.6  运输类发票自动核减说明', '9'),
        ('    3.7  核减发票上传', '9'),
        ('    3.8  查看报销状态', '10'),
        ('    3.9  处理被驳回的报销单', '10'),
        ('    3.10 确认收款', '11'),
        ('四、管理员操作指南', '11'),
        ('    4.1  审批中心概览', '11'),
        ('    4.2  审批报销单', '12'),
        ('    4.3  发起付款', '13'),
        ('    4.4  上传付款回单', '13'),
        ('    4.5  数据查询与发票管理', '14'),
        ('五、总经理操作指南', '15'),
        ('    5.1  审批商务报销', '15'),
        ('六、附录', '16'),
        ('    6.1  报销状态说明', '16'),
        ('    6.2  常见问题', '16'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run_item = p.add_run(item)
        set_font(run_item, size=11)
        # 填充点
        dots = '.' * max(1, 60 - len(item) - len(page))
        run_dots = p.add_run(dots)
        set_font(run_dots, size=11, color=(180,180,180))
        run_page = p.add_run(page)
        set_font(run_page, size=11)
    doc.add_page_break()

    # ===== 第一章：系统简介 =====
    add_h1(doc, '一、系统简介')
    add_para(doc,
        'YuliLog 财务区是公司 OA 工作日志系统的报销管理模块，实现了报销申请、'
        '审批、付款、确认收款的全流程数字化管理。员工无需纸质单据，所有操作均在'
        '系统中完成，审批进度实时可查。')
    doc.add_paragraph()
    add_para(doc, '系统支持三种报销类型：', bold=True)
    add_bullet(doc, '基础报销：金额小于 1000 元的日常费用（办公、差旅、服务等）')
    add_bullet(doc, '大额报销：金额大于等于 1000 元的费用，品类与基础报销相同')
    add_bullet(doc, '商务报销：商务接待、客户拜访、商务差旅、会议活动，不限金额，由总经理审批')
    doc.add_paragraph()
    add_para(doc, '系统角色说明：', bold=True)
    add_bullet(doc, '普通员工：提交报销单、上传发票、确认收款')
    add_bullet(doc, '管理员：审批基础/大额报销、处理付款、上传付款回单')
    add_bullet(doc, '总经理：审批商务报销')
    doc.add_page_break()

    # ===== 第二章：登录与账号设置 =====
    add_h1(doc, '二、登录与账号设置')

    add_h2(doc, '2.1  系统访问地址')
    data = [
        ['阶段', '访问地址', '使用条件'],
        ['当前内网阶段', 'http://192.168.100.7:8899', '需连接公司 Wi-Fi 或办公室内网'],
        ['正式上线后', 'https://yuliseek.com', '任意网络均可访问'],
    ]
    t = doc.add_table(rows=3, cols=3)
    t.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=10, bold=(i==0))
    doc.add_paragraph()
    add_warning(doc, '内网阶段请确保手机或电脑已连接公司 Wi-Fi，否则页面无法打开。推荐使用 Chrome 或 Edge 浏览器。')

    add_h2(doc, '2.2  登录步骤')
    add_para(doc, '账号由管理员统一创建，员工不能自行注册。管理员会将您的用户名和初始密码单独告知您。')
    doc.add_paragraph()
    add_step(doc, 1, '打开浏览器，在地址栏输入系统地址，按回车键进入登录页面。')
    add_step(doc, 2, '在"用户名"输入框填写您的账号（由管理员分配）。')
    add_step(doc, 3, '在"密码"输入框填写密码（至少6位，输入时显示为圆点属正常现象）。')
    add_step(doc, 4, '点击"登录"按钮，或直接按键盘 Enter 键。')
    add_step(doc, 5, '登录成功后，页面自动跳转到系统首页。')
    doc.add_paragraph()
    add_img(doc, 'login', '图1：系统登录页面')
    add_img(doc, 'home', '图2：登录成功后的系统首页')
    add_tip(doc, '首次登录后建议修改初始密码，点击右上角头像 → 用户设置 → 账号操作。')

    add_h2(doc, '2.3  完善个人信息（必做）')
    add_warning(doc, '使用报销功能前，必须先完善个人信息中的银行账号，否则财务付款时系统无法验证，报销流程将无法完成。')
    doc.add_paragraph()
    add_step(doc, 1, '登录后，点击左侧菜单"入职管理"，进入个人信息填写页面。')
    add_step(doc, 2, '填写基本信息：姓名、性别、出生日期、身份证号、联系方式等。')
    add_step(doc, 3, '找到页面下方的"收款信息"区域，填写以下四项（必填）：')
    add_bullet(doc, '收款人姓名：填写银行卡开户姓名（与银行卡一致）', indent_level=2)
    add_bullet(doc, '收款人手机：填写银行预留手机号', indent_level=2)
    add_bullet(doc, '开户行：填写银行名称，如"中国工商银行"', indent_level=2)
    add_bullet(doc, '银行卡号：填写完整的银行卡号', indent_level=2)
    add_step(doc, 4, '点击"保存草稿"保存信息，确认无误后点击"提交"。')
    doc.add_paragraph()
    add_img(doc, 'onboarding', '图3：入职信息页面（含收款信息填写区域）')
    add_tip(doc, '银行卡号请仔细核对，填错会导致付款失败，需联系管理员修改。')
    doc.add_page_break()

    # ===== 第三章：员工操作指南 =====
    add_h1(doc, '三、员工操作指南')

    add_h2(doc, '3.1  报销类型说明')
    data2 = [
        ['报销类型', '适用金额', '审批人', '典型场景'],
        ['基础报销', '< 1000 元', '管理员', '办公耗材、快递、差旅、茶水等'],
        ['大额报销', '≥ 1000 元', '管理员', '设备采购、大额服务费等'],
        ['商务报销', '不限金额', '总经理', '商务接待、客户拜访、会议活动等'],
    ]
    t2 = doc.add_table(rows=4, cols=4)
    t2.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(data2):
        for j, text in enumerate(row_data):
            cell = t2.rows[i].cells[j]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=10, bold=(i==0))
    doc.add_paragraph()
    add_tip(doc, '不确定选哪种类型时，看金额：< 1000元选基础报销，≥ 1000元选大额报销，商务活动选商务报销。')

    add_h2(doc, '3.2  创建基础报销单')
    add_para(doc, '适用于金额小于 1000 元的日常费用报销。')
    doc.add_paragraph()
    add_step(doc, 1, '点击左侧菜单"财务区" → "基础报销"，进入报销列表页面。')
    add_img(doc, 'reimb_list', '图4：基础报销列表页面（可查看所有报销单及状态）')
    add_step(doc, 2, '点击右上角"新建报销单"按钮，进入创建页面。')
    add_step(doc, 3, '填写报销基本信息：')
    add_bullet(doc, '报销标题：简要描述本次报销内容，如"2026年4月办公耗材采购"', indent_level=2)
    add_bullet(doc, '报销月份：选择费用发生的月份', indent_level=2)
    add_bullet(doc, '报销类别：从下拉菜单选择对应品类（如办公耗材、差旅费等）', indent_level=2)
    add_bullet(doc, '服务对象：填写服务的项目或部门名称', indent_level=2)
    add_img(doc, 'new_basic', '图5：新建基础报销单页面')
    add_step(doc, 4, '上传发票（详见 3.5 节），系统自动识别发票金额。')
    add_step(doc, 5, '确认发票信息无误后，点击"提交审批"正式提交；如需暂存，点击"保存草稿"。')
    add_tip(doc, '草稿状态下可随时修改，提交后进入审批流程，不可再修改。')

    add_h2(doc, '3.3  创建大额报销单')
    add_para(doc, '适用于金额大于等于 1000 元的费用报销，操作流程与基础报销完全相同。')
    add_step(doc, 1, '点击左侧菜单"财务区" → "大额报销"，点击"新建报销单"。')
    add_step(doc, 2, '填写报销信息，上传发票，提交审批。')
    add_img(doc, 'large_empty', '图6：新建大额报销单页面')
    add_img(doc, 'large_filled', '图7：已上传发票的大额报销单（系统自动汇总金额）')
    add_warning(doc, '大额报销单金额必须 ≥ 1000 元，否则请使用基础报销。')

    add_h2(doc, '3.4  创建商务报销单')
    add_para(doc, '适用于商务接待、客户拜访、商务差旅、会议活动等场景，由总经理审批。')
    add_step(doc, 1, '点击左侧菜单"财务区" → "商务报销"，点击"新建报销单"。')
    add_step(doc, 2, '除基础信息外，还需填写商务报销特有字段：')
    add_bullet(doc, '商务类型：选择商务接待 / 客户拜访 / 商务差旅 / 会议活动', indent_level=2)
    add_bullet(doc, '客户名称：填写对应的客户或合作方名称', indent_level=2)
    add_img(doc, 'business_empty', '图8：新建商务报销单页面（含商务类型和客户名称字段）')
    add_img(doc, 'business_filled', '图9：已填写完整的商务报销单')
    add_step(doc, 3, '上传发票后，点击"提交审批"，报销单将发送给总经理审批。')
    add_tip(doc, '商务报销由总经理审批，审批通过后由管理员处理付款，流程与其他类型略有不同。')
    doc.add_page_break()

    add_h2(doc, '3.5  上传发票与 OCR 识别')
    add_para(doc, '系统支持上传三种类型的附件：')
    add_bullet(doc, '普通发票：正式报销凭证，计入报销金额')
    add_bullet(doc, '无票收据：无正式发票时使用，计入报销金额')
    add_bullet(doc, '核减发票：不计入报销金额，仅作存档凭证（详见 3.7 节）')
    doc.add_paragraph()
    add_step(doc, 1, '在报销单详情页，点击"上传发票"按钮，选择发票文件（支持 PDF、PNG、JPG 格式，单文件不超过 10MB）。')
    add_step(doc, 2, '上传后，系统自动进行 OCR 识别，提取发票中的金额、日期、发票号码等信息。')
    add_img(doc, 'ocr_result', '图10：发票上传后 OCR 自动识别结果（金额、日期等自动填入）')
    add_step(doc, 3, '核对识别结果是否正确，如有误可手动修改。')
    add_step(doc, 4, '可重复上传多张发票，系统自动累加金额。')
    add_tip(doc, '发票图片请保持清晰，模糊图片可能导致 OCR 识别失败，需手动填写金额。')
    add_warning(doc, '发票金额以 OCR 识别结果为准，提交前请务必核对总金额是否正确。')

    add_h2(doc, '3.6  运输类发票自动核减说明')
    add_para(doc,
        '系统对运输、交通、汽油、柴油类发票设有月度额度限制：'
        '每人每月最多报销 1500 元。超出部分系统将自动核减，无需手动操作。')
    doc.add_paragraph()
    add_img(doc, 'auto_deduction', '图11：运输类发票超出月度额度后，系统自动核减示例（本月已超出，自动扣减超出部分）')
    add_para(doc, '核减规则：', bold=True)
    add_bullet(doc, '核减金额 = 本月运输类发票总额 - 1500 元')
    add_bullet(doc, '核减金额按比例分配到每张运输类发票上')
    add_bullet(doc, '实际报销金额 = 发票总额 - 核减金额')
    add_tip(doc, '系统会在发票列表中显示每张发票的核减金额，提交前可查看。')

    add_h2(doc, '3.7  核减发票上传')
    add_para(doc,
        '核减发票是指员工持有但不能报销的发票（如超标费用、个人消费等），'
        '上传后仅作存档凭证，不计入报销金额。')
    add_step(doc, 1, '在报销单详情页，点击"上传核减发票"按钮。')
    add_step(doc, 2, '选择对应的发票文件上传，系统标记为核减，不计入报销总额。')
    add_img(doc, 'manual_deduction', '图12：核减发票上传后的显示效果（标注为核减，不计入金额）')
    add_tip(doc, '核减发票与普通发票分开显示，报销总额不包含核减发票金额。')

    add_h2(doc, '3.8  查看报销状态')
    add_step(doc, 1, '在报销列表页，每条报销单右侧显示当前状态标签。')
    add_step(doc, 2, '点击任意报销单可查看详情，包括审批进度、发票列表、付款信息等。')
    add_img(doc, 'submitted_list', '图13：报销列表页面，显示各报销单的当前状态')
    doc.add_paragraph()
    add_status_table(doc)

    add_h2(doc, '3.9  处理被驳回的报销单')
    add_para(doc, '审批人驳回报销单后，您会收到通知，可查看驳回原因并修改后重新提交。')
    add_step(doc, 1, '在报销列表中找到状态为"已驳回"的报销单，点击查看详情。')
    add_step(doc, 2, '查看审批流程中的驳回原因（审批人会填写具体原因）。')
    add_img(doc, 'rejected', '图14：报销单被驳回后的审批流程弹窗，显示驳回原因')
    add_step(doc, 3, '根据驳回原因修改报销单内容（如更换发票、修改金额等）。')
    add_step(doc, 4, '修改完成后，重新点击"提交审批"。')
    add_tip(doc, '驳回后的报销单可以修改所有内容，包括发票、金额、报销类别等。')

    add_h2(doc, '3.10  确认收款')
    add_para(doc, '管理员完成付款并上传付款回单后，报销单状态变为"待确认收款"，需要您确认收到款项。')
    add_step(doc, 1, '在报销列表中找到状态为"待确认收款"的报销单，点击查看详情。')
    add_step(doc, 2, '核对付款金额是否与报销单一致。')
    add_img(doc, 'pay_uploaded', '图15：待确认收款状态，显示付款回单信息')
    add_step(doc, 3, '确认收到款项后，点击"确认收款"按钮，报销单状态变为"已完成"。')
    add_img(doc, 'completed_list', '图16：已完成的报销单列表')
    add_warning(doc, '请在收到款项后再点击确认，确认后流程结束，不可撤销。')
    doc.add_page_break()

    # ===== 第四章：管理员操作指南 =====
    add_h1(doc, '四、管理员操作指南')
    add_para(doc, '管理员负责审批基础报销和大额报销，并处理所有类型报销的付款操作。')

    add_h2(doc, '4.1  审批中心概览')
    add_step(doc, 1, '点击左侧菜单"审批中心"，进入管理员审批工作台。')
    add_step(doc, 2, '页面顶部显示待处理数量统计，包括待审批、待付款等各类待办事项数量。')
    add_step(doc, 3, '列表中显示所有待审批的报销单，包含提交人、报销类型、金额、提交时间等信息。')
    add_img(doc, 'admin_approval', '图17：管理员审批中心，显示待审批报销单列表及统计数据')

    add_h2(doc, '4.2  审批报销单')
    add_step(doc, 1, '在审批中心列表中，点击某条报销单的"详情"按钮，查看完整报销信息。')
    add_step(doc, 2, '仔细核对以下内容：')
    add_bullet(doc, '报销类别是否正确', indent_level=2)
    add_bullet(doc, '发票是否清晰、真实', indent_level=2)
    add_bullet(doc, '报销金额是否合理', indent_level=2)
    add_bullet(doc, '是否符合公司报销政策', indent_level=2)
    add_step(doc, 3, '审批通过：点击"通过"按钮，报销单进入待付款状态。')
    add_step(doc, 4, '审批驳回：点击"驳回"按钮，在弹窗中填写驳回原因，员工将收到通知并可修改重提。')
    add_img(doc, 'admin_flow', '图18：审批流程弹窗，显示审批历史记录（含驳回原因和重新提交记录）')
    add_tip(doc, '驳回时请填写具体原因，帮助员工了解需要修改的内容，避免反复驳回。')
    add_warning(doc, '管理员只能审批基础报销和大额报销，商务报销由总经理审批。')

    add_h2(doc, '4.3  发起付款')
    add_para(doc, '报销单审批通过后，进入"已批准待付款"列表，管理员需完成付款操作。')
    doc.add_paragraph()
    add_h3(doc, '单笔付款')
    add_step(doc, 1, '在"已批准待付款"列表中找到对应报销单，点击"发起付款"。')
    add_step(doc, 2, '确认收款人姓名、银行账号、付款金额无误。')
    add_step(doc, 3, '通过银行系统完成转账操作。')
    add_img(doc, 'admin_pay_detail', '图19：付款详情页，显示收款人信息和报销金额')

    add_h3(doc, '批量付款')
    add_step(doc, 1, '在"已批准待付款"列表中，勾选多条报销单（需同一收款人、同一报销类型）。')
    add_step(doc, 2, '点击"批量付款"，系统自动合并金额，一次转账完成多笔报销。')
    add_img(doc, 'admin_pending', '图20：待付款列表，勾选多条记录进行批量付款操作')
    add_tip(doc, '批量付款要求：同一收款人 + 同一报销类型，可大幅减少转账次数。')

    add_h2(doc, '4.4  上传付款回单')
    add_para(doc, '完成银行转账后，需将付款回单（银行转账截图）上传到系统，系统自动 OCR 验证。')
    add_step(doc, 1, '在付款详情页，点击"上传付款回单"按钮。')
    add_step(doc, 2, '选择付款回单图片（银行 App 转账成功截图），上传到系统。')
    add_img(doc, 'upload_proof', '图21：上传付款回单页面，支持拖拽上传或点击选择文件')
    add_step(doc, 3, '系统自动 OCR 识别回单中的收款人姓名、银行账号、付款金额，并与报销单信息进行比对验证。')
    add_step(doc, 4, '验证通过后，点击确认，报销单状态变为"待确认收款"，等待员工确认。')
    add_img(doc, 'proof_done', '图22：付款回单上传成功，等待员工确认收款')
    add_warning(doc, '付款回单中的收款账号必须与员工填写的银行账号完全一致，否则 OCR 验证不通过。如验证失败，请检查员工个人信息中的银行账号是否填写正确。')

    add_h2(doc, '4.5  数据查询与发票管理')
    add_step(doc, 1, '在审批中心点击"全部查询"标签，可查看所有报销单的汇总统计数据。')
    add_step(doc, 2, '支持按报销类型、状态、时间范围等条件筛选，并可导出数据。')
    add_img(doc, 'admin_query', '图23：全部查询页面，显示报销汇总统计和筛选功能')
    add_step(doc, 3, '点击"发票管理"标签，可查看所有上传的发票记录，支持按条件搜索。')
    add_img(doc, 'admin_invoice', '图24：发票管理页面，可查询和管理所有发票记录')
    doc.add_page_break()

    # ===== 第五章：总经理操作指南 =====
    add_h1(doc, '五、总经理操作指南')
    add_para(doc, '总经理负责审批所有商务报销单，基础报销和大额报销不在总经理审批范围内。')

    add_h2(doc, '5.1  审批商务报销')
    add_step(doc, 1, '登录系统后，点击左侧菜单"总经理审批中心"。')
    add_img(doc, 'gm_home', '图25：总经理登录后的系统首页')
    add_step(doc, 2, '审批中心显示所有待审批的商务报销单，包含提交人、商务类型、金额等信息。')
    add_img(doc, 'gm_approval', '图26：总经理审批中心，显示待审批商务报销列表')
    add_step(doc, 3, '点击"详情"查看报销单完整信息，重点核对：')
    add_bullet(doc, '商务类型是否与实际情况相符', indent_level=2)
    add_bullet(doc, '客户名称是否正确', indent_level=2)
    add_bullet(doc, '发票金额是否合理', indent_level=2)
    add_img(doc, 'gm_detail', '图27：商务报销审批详情页，显示商务类型、客户名称及发票信息')
    add_step(doc, 4, '审批通过：点击"通过"按钮，报销单转交管理员处理付款。')
    add_step(doc, 5, '审批驳回：点击"驳回"按钮，填写驳回原因，员工收到通知后可修改重提。')
    add_tip(doc, '商务报销审批通过后，后续付款由管理员操作，总经理无需再处理。')
    doc.add_page_break()

    # ===== 第六章：附录 =====
    add_h1(doc, '六、附录')

    add_h2(doc, '6.1  报销状态说明')
    add_status_table(doc)

    add_h2(doc, '6.2  常见问题')
    faqs = [
        ('登录时提示"用户名或密码错误"怎么办？',
         '请确认用户名和密码是否正确，注意区分大小写。如忘记密码，请联系管理员重置。'),
        ('页面打不开怎么办？',
         '内网阶段请确认已连接公司 Wi-Fi。如已连接仍无法访问，请尝试刷新页面或更换 Chrome/Edge 浏览器。'),
        ('发票上传后 OCR 识别金额有误怎么办？',
         '可手动修改识别结果中的金额，确认无误后再提交。'),
        ('运输类发票被自动核减了，能申诉吗？',
         '运输类发票月度额度为 1500 元，超出部分按公司政策自动核减，如有特殊情况请联系管理员说明。'),
        ('报销单提交后能修改吗？',
         '提交后进入审批流程，不可修改。如需修改，可联系审批人驳回后再修改重提。'),
        ('付款回单 OCR 验证不通过怎么办？',
         '请检查员工个人信息中的银行账号是否填写正确，账号需与付款回单中的收款账号完全一致。'),
        ('忘记确认收款怎么办？',
         '报销单状态为"待确认收款"时，随时可以登录系统点击确认收款，没有时间限制。'),
        ('手机能使用系统吗？',
         '可以，系统支持手机浏览器访问，使用 Chrome 或 Safari 浏览器打开系统地址即可。'),
    ]
    for q, a in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.left_indent = Cm(0.3)
        run_q = p_q.add_run(f'问：{q}')
        set_font(run_q, size=11, bold=True)
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Cm(0.75)
        run_a = p_a.add_run(f'答：{a}')
        set_font(run_a, size=11, color=(80, 80, 80))
        doc.add_paragraph()

    add_h2(doc, '6.3  联系支持')
    add_para(doc, '如在使用过程中遇到系统问题，请联系：')
    add_bullet(doc, '系统管理员：羽隶（微信 / 电话）')
    add_bullet(doc, '反馈渠道：在公司微信群中说明问题，1个工作日内响应')

    # ===== 保存 =====
    doc.save(OUTPUT_PATH)
    print(f'✅ 文档已生成：{OUTPUT_PATH}')


if __name__ == '__main__':
    create_manual()
